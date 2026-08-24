"""File content extraction for uploaded documents and images."""

import csv
import io
import os
from typing import Optional

from app.core.config import settings
from app.services.llm_service import LLMImage, LLMMessage, encode_image_bytes, get_llm_provider

IMAGE_EXTS = {"png", "jpg", "jpeg", "webp", "gif"}
ALLOWED_EXTS = IMAGE_EXTS | {"pdf", "docx", "xlsx", "csv", "txt"}


def is_allowed(filename: str) -> bool:
    return ext_of(filename) in ALLOWED_EXTS


def is_image(filename: str) -> bool:
    return ext_of(filename) in IMAGE_EXTS


def ext_of(filename: str) -> str:
    return os.path.splitext(filename)[1].lower().lstrip(".")


def sniff_media_type(data: bytes) -> Optional[str]:
    """Detect image media type from magic bytes (uploads may lack proper extensions)."""
    if data[:8] == b"\x89PNG\r\n\x1a\n":
        return "image/png"
    if data[:3] == b"\xff\xd8\xff":
        return "image/jpeg"
    if data[:4] == b"RIFF" and data[8:12] == b"WEBP":
        return "image/webp"
    if data[:6] in (b"GIF87a", b"GIF89a"):
        return "image/gif"
    return None


def extract_pdf(raw: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(raw))
    pages = []
    for i, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(f"[Page {i + 1}]\n{text}")
    return "\n\n".join(pages)


def extract_docx(raw: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(raw))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_xlsx(raw: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
    sheets = []
    for ws in wb.worksheets:
        rows = []
        for row in ws.iter_rows(values_only=True):
            values = ["" if v is None else str(v) for v in row]
            if any(v.strip() for v in values):
                rows.append(", ".join(values))
        if rows:
            sheets.append(f"### Sheet: {ws.title}\n" + "\n".join(rows))
    return "\n\n".join(sheets)


def extract_csv(raw: bytes) -> str:
    text = raw.decode("utf-8-sig", errors="replace")
    rows = list(csv.reader(io.StringIO(text)))
    lines = []
    for row in rows[:2000]:  # cap to keep prompts sane
        lines.append(", ".join(cell.strip() for cell in row))
    return "\n".join(lines)


def extract_text_file(raw: bytes) -> str:
    return raw.decode("utf-8", errors="replace")


def extract_document(filename: str, raw: bytes) -> str:
    """Extract text content from a non-image upload."""
    ext = ext_of(filename)
    try:
        if ext == "pdf":
            return extract_pdf(raw)
        if ext == "docx":
            return extract_docx(raw)
        if ext == "xlsx":
            return extract_xlsx(raw)
        if ext == "csv":
            return extract_csv(raw)
        return extract_text_file(raw)
    except Exception as exc:
        return f"(Could not extract content from '{filename}': {exc})"


async def describe_image(raw: bytes, filename: str) -> str:
    """Use a vision-capable LLM call to describe/extract the image content."""
    media = sniff_media_type(raw) or f"image/{ext_of(filename)}"
    provider = get_llm_provider()
    prompt = (
        "Analyze this image thoroughly. Extract all visible text verbatim (OCR), then "
        "describe charts/diagrams/photos precisely so another AI can answer questions "
        "about it without seeing it."
    )
    try:
        chunks = []
        async for chunk in provider.stream_generate(
            [
                LLMMessage(
                    role="user",
                    content=prompt,
                    images=[LLMImage(media_type=media, data_b64=encode_image_bytes(raw))],
                )
            ]
        ):
            chunks.append(chunk)
        result = "".join(chunks).strip()
        return result or "(No description produced.)"
    except Exception as exc:
        return f"(Image analysis unavailable: {exc})"


def save_upload(filename: str, raw: bytes) -> str:
    """Persist an upload under UPLOAD_DIR with a unique name; returns relative path."""
    safe_name = os.path.basename(filename).replace(os.sep, "_")[-150:]
    import uuid

    stored = f"{uuid.uuid4().hex}_{safe_name}"
    path = os.path.join(settings.upload_dir, stored)
    with open(path, "wb") as fh:
        fh.write(raw)
    return stored
